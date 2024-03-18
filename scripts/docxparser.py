from docx import Document
import json
doc = Document('pie.docx')

# print(doc.paragraphs)
data = []
for paragraph in doc.paragraphs:
    data.append(paragraph.text)

json_data = json.dumps(data, indent=4)
print(json_data)